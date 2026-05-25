"""
Text Extractor — extracts plain text from PDF, DOCX, RTF, TXT bytes.
Multiple strategies with fallbacks to maximize extraction success.
"""
import io
import re
import struct
import zipfile
from typing import Optional


def detect_file_type(data: bytes) -> str:
    """Detect file type from magic bytes."""
    if not data or len(data) < 8:
        return "unknown"

    # PDF: starts with %PDF
    if data[:4] == b"%PDF":
        return "pdf"

    # DOCX/XLSX/PPTX: ZIP-based Office Open XML
    if data[:4] == b"PK\x03\x04":
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                names = zf.namelist()
                if any("word/" in n for n in names):
                    return "docx"
                if any("xl/" in n for n in names):
                    return "xlsx"
                if any("ppt/" in n for n in names):
                    return "pptx"
        except:
            pass
        return "docx"  # Assume docx for Office ZIP

    # DOC (old binary format): starts with D0 CF 11 E0
    if data[:4] == b"\xd0\xcf\x11\xe0":
        return "doc"

    # RTF: starts with {\rtf
    if data[:5] == b"{\\rtf":
        return "rtf"

    # Plain text heuristic: check if mostly printable ASCII/UTF-8
    try:
        sample = data[:2000].decode("utf-8")
        printable_ratio = sum(1 for c in sample if c.isprintable() or c.isspace()) / len(sample)
        if printable_ratio > 0.85:
            return "txt"
    except:
        pass

    # HTML (some resume links return HTML pages)
    if b"<html" in data[:500].lower() or b"<!doctype" in data[:500].lower():
        return "html"

    return "unknown"


def extract_text_from_bytes(data: bytes, file_type: str) -> str:
    """Extract text using the appropriate parser for the file type."""
    extractors = {
        "pdf": [_extract_pdf_pdfplumber, _extract_pdf_fallback],
        "docx": [_extract_docx, _extract_docx_fallback],
        "doc": [_extract_doc_fallback],
        "rtf": [_extract_rtf],
        "txt": [_extract_txt],
        "html": [_extract_html],
        "unknown": [_extract_pdf_pdfplumber, _extract_docx, _extract_txt, _extract_html],
    }

    strategies = extractors.get(file_type, extractors["unknown"])

    for strategy in strategies:
        try:
            text = strategy(data)
            if text and len(text.strip()) > 10:
                return _clean_text(text)
        except Exception:
            continue

    return ""


def _extract_pdf_pdfplumber(data: bytes) -> str:
    """Primary PDF extractor using pdfplumber (handles most PDFs well)."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages[:20]:  # Limit to 20 pages
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

            # Also try extracting from tables
            tables = page.extract_tables()
            for table in (tables or []):
                for row in (table or []):
                    cells = [str(c).strip() for c in (row or []) if c]
                    if cells:
                        text_parts.append(" | ".join(cells))

    return "\n".join(text_parts)


def _extract_pdf_fallback(data: bytes) -> str:
    """Fallback PDF extraction using raw stream parsing."""
    text = ""
    # Try to extract text streams from raw PDF
    try:
        content = data.decode("latin-1")
        # Find text between BT and ET markers
        bt_et_pattern = re.findall(r"BT\s*(.*?)\s*ET", content, re.DOTALL)
        for block in bt_et_pattern:
            # Extract text from Tj and TJ operators
            tj_matches = re.findall(r"\((.*?)\)\s*Tj", block)
            text += " ".join(tj_matches) + "\n"

            # TJ array
            tj_array = re.findall(r"\[(.*?)\]\s*TJ", block)
            for arr in tj_array:
                parts = re.findall(r"\((.*?)\)", arr)
                text += "".join(parts) + " "
    except:
        pass

    return text


def _extract_docx(data: bytes) -> str:
    """Primary DOCX extractor using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []

    # Main body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Tables (resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    # Headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())

    return "\n".join(parts)


def _extract_docx_fallback(data: bytes) -> str:
    """Fallback DOCX extraction by directly reading XML from ZIP."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            # Read main document XML
            if "word/document.xml" in zf.namelist():
                xml_content = zf.read("word/document.xml").decode("utf-8")
                # Strip XML tags to get text
                text = re.sub(r"<[^>]+>", " ", xml_content)
                text = re.sub(r"\s+", " ", text)
                return text.strip()
    except:
        pass
    return ""


def _extract_doc_fallback(data: bytes) -> str:
    """Extract text from old .doc format using binary parsing."""
    # Simple extraction: look for readable ASCII/Unicode strings
    text_parts = []
    try:
        # Try UTF-16 LE strings (common in .doc)
        i = 0
        while i < len(data) - 1:
            if 32 <= data[i] < 127 and data[i + 1] == 0:
                # Possible UTF-16 LE string
                end = i
                chars = []
                while end < len(data) - 1 and 32 <= data[end] < 127 and data[end + 1] == 0:
                    chars.append(chr(data[end]))
                    end += 2
                if len(chars) >= 4:
                    text_parts.append("".join(chars))
                i = end
            else:
                i += 1
    except:
        pass

    if not text_parts:
        # Fallback: extract ASCII strings
        ascii_strings = re.findall(rb"[\x20-\x7e]{4,}", data)
        text_parts = [s.decode("ascii") for s in ascii_strings[:500]]

    return "\n".join(text_parts)


def _extract_rtf(data: bytes) -> str:
    """Extract text from RTF format."""
    try:
        content = data.decode("latin-1")
        # Remove RTF control words and groups
        text = re.sub(r"\\[a-z]+\d*\s?", " ", content)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()
    except:
        return ""


def _extract_txt(data: bytes) -> str:
    """Extract text from plain text files."""
    # Try multiple encodings
    for encoding in ["utf-8", "utf-16", "latin-1", "cp1252", "iso-8859-1"]:
        try:
            return data.decode(encoding)
        except:
            continue
    return ""


def _extract_html(data: bytes) -> str:
    """Extract text from HTML (some resume links return HTML pages)."""
    try:
        from bs4 import BeautifulSoup
        # Try decoding
        for encoding in ["utf-8", "latin-1"]:
            try:
                html = data.decode(encoding)
                break
            except:
                continue
        else:
            return ""

        soup = BeautifulSoup(html, "html.parser")

        # Remove script and style elements
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n")
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    except:
        return ""


def _clean_text(text: str) -> str:
    """Clean extracted text — normalize whitespace, remove garbage."""
    # Remove null bytes
    text = text.replace("\x00", "")
    # Normalize unicode whitespace
    text = re.sub(r"[​‌‍﻿]", "", text)
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove very long lines of dots/dashes (table of contents style)
    text = re.sub(r"[.\-_]{10,}", " ", text)
    return text.strip()
